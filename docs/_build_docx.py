"""Build a session results docx for an autoresearch session.

Reads:
  results/{date}_{scope}/README.md
  results/{date}_{scope}/iter-*/summary.md

Writes:
  results/{date}_{scope}/SESSION_REPORT.docx

Style: conservative, paper-like. Black body text on white. Brand-color accents
on H1 and H2 only (turquoise / blueviolet). Geist for headings, Geist Mono for
inline code and metric values. Slugs humanized for display. No internal
references — this is a publication-style artifact.

Run:
    python docs/_build_docx.py --date 2026-04-30 --scope fw-arch-sweep
"""
from __future__ import annotations

import argparse
import re
from datetime import datetime
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    raise SystemExit("python-docx not installed. Run: pip install python-docx")


# ----- Brand palette (used sparingly in docx) -----
TURQUOISE  = RGBColor(0x40, 0xE0, 0xD0)
DEEPPINK   = RGBColor(0xFF, 0x14, 0x93)
BLUEVIOLET = RGBColor(0x8A, 0x2B, 0xE2)
INK        = RGBColor(0x14, 0x14, 0x18)
MUTED      = RGBColor(0x55, 0x55, 0x5C)

MONO = "Geist Mono"
BODY = "Geist"


def _humanize(slug: str) -> str:
    if not slug:
        return ""
    slug = re.sub(r"^iter-\d+_", "", slug)
    parts = re.split(r"[-_]+", slug)
    return " ".join(p[:1].upper() + p[1:] if p else "" for p in parts).strip()


def _styled(p, text, *, font=BODY, size=11, color=INK, bold=False, italic=False):
    r = p.add_run(text)
    r.font.name = font
    r.font.size = Pt(size)
    r.font.color.rgb = color
    r.bold = bold
    r.italic = italic
    return r


def _add_horizontal_rule(doc):
    """Insert a thin horizontal rule (border-bottom on a paragraph)."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "DDDDDD")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def _render_md_table(doc, table_lines: list[str]) -> None:
    """Render a markdown table (header row + separator + N data rows) as a
    proper docx table. Skips the separator. Preserves cell text only — no
    inline-bold parsing inside cells (cheap and good enough for AUC tables)."""
    rows = []
    for ln in table_lines:
        if re.match(r"^\s*\|?\s*[-: |]+\|?\s*$", ln) and "|" in ln:
            # Separator line
            continue
        cells = [c.strip() for c in ln.strip().strip("|").split("|")]
        rows.append(cells)
    if not rows:
        return
    n_cols = max(len(r) for r in rows)
    rows = [r + [""] * (n_cols - len(r)) for r in rows]
    tbl = doc.add_table(rows=len(rows), cols=n_cols)
    tbl.style = "Light Grid Accent 1"
    for ri, row in enumerate(rows):
        for ci, cell_text in enumerate(row):
            cell = tbl.rows[ri].cells[ci]
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            _styled(p, cell_text, font=MONO, size=9,
                    color=MUTED if ri == 0 else INK,
                    bold=(ri == 0))


def _render_md_body(doc, body_text: str) -> None:
    """Walk a markdown body and emit docx paragraphs / tables.

    Recognized blocks:
      - h2/h3 headings → bold paragraph
      - tables (lines beginning + ending with `|`) → docx tables
      - bullets (`-` or `*` at start) → bulleted paragraph
      - fenced code blocks (```...```) → mono paragraph block
      - everything else → body paragraph (with **bold** spans bolded)
    """
    lines = body_text.splitlines()
    i = 0
    in_fence = False
    fence_buf: list[str] = []
    while i < len(lines):
        ln = lines[i]
        stripped = ln.strip()

        # Fenced code
        if stripped.startswith("```"):
            if not in_fence:
                in_fence = True
                fence_buf = []
            else:
                # Close fence — emit buffered lines as mono paragraphs
                for fl in fence_buf:
                    p = doc.add_paragraph()
                    p.paragraph_format.space_after = Pt(0)
                    _styled(p, fl, font=MONO, size=9, color=INK)
                in_fence = False
                fence_buf = []
            i += 1
            continue
        if in_fence:
            fence_buf.append(ln)
            i += 1
            continue

        # Heading
        m = re.match(r"^(#{2,4})\s+(.+)", stripped)
        if m:
            level = len(m.group(1))
            heading_text = m.group(2).strip()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(2)
            size = {2: 12, 3: 11, 4: 10}.get(level, 11)
            _styled(p, heading_text, font=BODY, size=size, color=INK, bold=True)
            i += 1
            continue

        # Table — collect contiguous lines that contain "|"
        if "|" in ln and ln.lstrip().startswith("|"):
            tbl_lines = [ln]
            j = i + 1
            while j < len(lines) and "|" in lines[j] and lines[j].lstrip().startswith("|"):
                tbl_lines.append(lines[j])
                j += 1
            _render_md_table(doc, tbl_lines)
            i = j
            continue

        # Bullet
        if stripped.startswith(("- ", "* ")):
            text = stripped[2:].strip()
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(0)
            _emit_inline(p, text)
            i += 1
            continue

        # Blank
        if not stripped:
            i += 1
            continue

        # Plain paragraph — collect until next blank / structural marker
        para_lines = [stripped]
        j = i + 1
        while j < len(lines):
            nxt = lines[j].strip()
            if (not nxt or nxt.startswith("#") or "|" in nxt
                or nxt.startswith(("- ", "* ", "```"))):
                break
            para_lines.append(nxt)
            j += 1
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        _emit_inline(p, " ".join(para_lines))
        i = j


def _emit_inline(p, text: str) -> None:
    """Emit `text` into paragraph `p`, bolding **bold** spans and rendering
    `code` spans in Geist Mono."""
    pos = 0
    pattern = re.compile(r"\*\*(.+?)\*\*|`([^`]+)`")
    for m in pattern.finditer(text):
        if m.start() > pos:
            _styled(p, text[pos:m.start()], font=BODY, size=10, color=INK)
        if m.group(1) is not None:
            _styled(p, m.group(1), font=BODY, size=10, color=INK, bold=True)
        else:
            _styled(p, m.group(2), font=MONO, size=9, color=INK)
        pos = m.end()
    if pos < len(text):
        _styled(p, text[pos:], font=BODY, size=10, color=INK)


def _read_session(date_str, scope):
    session_dir = Path("results") / f"{date_str}_{scope}"
    readme = session_dir / "README.md"
    if not readme.exists():
        raise SystemExit(f"no session at {session_dir}")
    text = readme.read_text()
    scope_text = re.search(r"\*\*Scope:\*\*\s*(.+)", text)
    target = re.search(r"\*\*Target metric:\*\*\s*(.+)", text)

    candidates = []
    for d in sorted(session_dir.glob("iter-*")):
        if not d.is_dir():
            continue
        m = re.match(r"iter-(\d+)_(.+)", d.name)
        if not m:
            continue
        summary = d / "summary.md"
        body = summary.read_text() if summary.exists() else ""
        metric_match = re.search(r"^\s*metric\s*[:|]\s*(.+)$", body, re.MULTILINE | re.IGNORECASE)
        status_match = re.search(r"^\s*status\s*[:|]\s*(.+)$", body, re.MULTILINE | re.IGNORECASE)
        figs = sorted(d.glob("fig_*.png"))
        candidates.append({
            "iter": int(m.group(1)),
            "candidate": m.group(2),
            "summary": body,
            "metric": metric_match.group(1).strip() if metric_match else "—",
            "status": status_match.group(1).strip() if status_match else "—",
            "figures": figs,
        })

    return {
        "scope_text": scope_text.group(1).strip() if scope_text else _humanize(scope),
        "target": target.group(1).strip() if target else "",
        "candidates": candidates,
    }


def main():
    p = argparse.ArgumentParser(description=__doc__)
    p.add_argument("--date", default=datetime.now().strftime("%Y-%m-%d"))
    p.add_argument("--scope", required=True)
    args = p.parse_args()

    path_date = args.date.replace("-", "")
    session = _read_session(path_date, args.scope)
    out_dir = Path("results") / f"{path_date}_{args.scope}"
    out_dir.mkdir(parents=True, exist_ok=True)
    out = out_dir / "SESSION_REPORT.docx"

    doc = Document()

    # Eyebrow + title
    eyebrow = doc.add_paragraph()
    _styled(eyebrow, "RESEARCH SESSION", font=MONO, size=9, color=MUTED, bold=True)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(2)
    _styled(title, _humanize(args.scope), font=BODY, size=24, color=TURQUOISE, bold=True)

    # Scope sentence + meta
    meta = doc.add_paragraph()
    meta.paragraph_format.space_after = Pt(4)
    _styled(meta, session["scope_text"], font=BODY, size=11, color=INK)

    if session["target"]:
        target_p = doc.add_paragraph()
        target_p.paragraph_format.space_after = Pt(2)
        _styled(target_p, "Target — ", font=MONO, size=10, color=MUTED, bold=True)
        _styled(target_p, session["target"], font=MONO, size=10, color=BLUEVIOLET, bold=True)

    date_p = doc.add_paragraph()
    _styled(date_p, args.date, font=MONO, size=9, color=MUTED)

    _add_horizontal_rule(doc)

    # Iterations
    h = doc.add_paragraph()
    h.paragraph_format.space_before = Pt(8)
    _styled(h, "Iterations", font=BODY, size=14, color=DEEPPINK, bold=True)

    for c in session["candidates"]:
        # Section header: candidate name
        sec = doc.add_paragraph()
        sec.paragraph_format.space_before = Pt(10)
        sec.paragraph_format.space_after = Pt(2)
        _styled(sec, _humanize(c["candidate"]), font=BODY, size=12, color=INK, bold=True)

        # Status + metric line
        meta_line = doc.add_paragraph()
        meta_line.paragraph_format.space_after = Pt(4)
        _styled(meta_line, "Status: ", font=MONO, size=9, color=MUTED, bold=True)
        _styled(meta_line, c["status"], font=MONO, size=9, color=INK)
        _styled(meta_line, "      Metric: ", font=MONO, size=9, color=MUTED, bold=True)
        _styled(meta_line, str(c["metric"]), font=MONO, size=9, color=BLUEVIOLET, bold=True)

        # Body — render markdown structurally: headings as bold lines,
        # tables as docx tables, bullets as bullet paragraphs, fences as
        # mono blocks. Strip the leading H1 (iter-block already shows the
        # candidate name).
        body_text = (c["summary"] or "").strip()
        body_text = re.sub(r"\A#\s+[^\n]*\n+", "", body_text)
        if body_text:
            _render_md_body(doc, body_text)

        # Figures — embed every fig_*.png in the iter dir, page-width-ish.
        for fig in c.get("figures", []):
            try:
                doc.add_picture(str(fig), width=Inches(6.0))
                cap = doc.add_paragraph()
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cap.paragraph_format.space_after = Pt(8)
                _styled(cap, fig.stem, font=MONO, size=8, color=MUTED, italic=True)
            except Exception as exc:
                # Skip unreadable / oversized images rather than aborting.
                err = doc.add_paragraph()
                _styled(err, f"[figure unavailable: {fig.name} ({exc})]",
                        font=MONO, size=8, color=MUTED, italic=True)

    _add_horizontal_rule(doc)

    # Results table
    h2 = doc.add_paragraph()
    h2.paragraph_format.space_before = Pt(10)
    _styled(h2, "Results", font=BODY, size=14, color=DEEPPINK, bold=True)

    if session["candidates"]:
        tbl = doc.add_table(rows=1 + len(session["candidates"]), cols=4)
        tbl.style = "Light List Accent 1"
        hdr = tbl.rows[0].cells
        for i, label in enumerate(("Iter", "Candidate", "Status", "Metric")):
            hdr_p = hdr[i].paragraphs[0]
            _styled(hdr_p, label, font=MONO, size=9, color=MUTED, bold=True)
        for row_i, c in enumerate(session["candidates"]):
            cells = tbl.rows[row_i + 1].cells
            _styled(cells[0].paragraphs[0], f"{c['iter']:02d}",
                    font=MONO, size=9, color=INK)
            _styled(cells[1].paragraphs[0], _humanize(c["candidate"]),
                    font=BODY, size=10, color=INK)
            _styled(cells[2].paragraphs[0], c["status"],
                    font=MONO, size=9, color=INK)
            _styled(cells[3].paragraphs[0], str(c["metric"]),
                    font=MONO, size=9, color=BLUEVIOLET, bold=True)

    doc.save(out)
    print(f"wrote {out}")


if __name__ == "__main__":
    main()
